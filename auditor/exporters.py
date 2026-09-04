"""Download formats for the feedback document: .txt, .docx and .pdf."""

from __future__ import annotations

import io

from .feedback import render_text


def to_txt(document: dict) -> bytes:
    return render_text(document).encode("utf-8")


def to_docx(document: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    heading = doc.add_paragraph()
    run = heading.add_run(document["title"])
    run.bold = True
    run.font.size = Pt(15)

    date_line = doc.add_paragraph()
    date_line.add_run(f"Date: {document['date']}")

    if document["clean"]:
        doc.add_paragraph("No data-quality issues were found in this workbook. "
                          "No corrections are needed.")
    else:
        for section in document["sections"]:
            title = doc.add_paragraph()
            title_run = title.add_run(f"{section['number']}. {section['title']}")
            title_run.bold = True
            title_run.font.size = Pt(12)

            doc.add_paragraph(section["problem"])

            please = doc.add_paragraph()
            please.add_run("Please:")
            for action in section["actions"]:
                doc.add_paragraph(action, style="List Bullet")

            for detail in section["details"]:
                note = doc.add_paragraph()
                note.add_run(detail["heading"]).italic = True
                for item in detail["items"]:
                    doc.add_paragraph(item, style="List Bullet")
                if detail["more"]:
                    doc.add_paragraph(f"... and {detail['more']} more",
                                      style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def to_pdf(document: dict) -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    def escape(text):
        return (str(text).replace("&", "&amp;")
                .replace("<", "&lt;").replace(">", "&gt;"))

    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Helvetica-Bold",
                                fontSize=15, alignment=TA_LEFT, spaceAfter=4),
        "meta": ParagraphStyle("meta", parent=base["Normal"], fontName="Helvetica",
                               fontSize=10, textColor="#555555", spaceAfter=14),
        "heading": ParagraphStyle("heading", parent=base["Normal"], fontName="Helvetica-Bold",
                                  fontSize=11.5, spaceBefore=12, spaceAfter=5),
        "body": ParagraphStyle("body", parent=base["Normal"], fontName="Helvetica",
                               fontSize=10.5, leading=15, spaceAfter=6),
        "note": ParagraphStyle("note", parent=base["Normal"], fontName="Helvetica-Oblique",
                               fontSize=9.5, leading=14, spaceBefore=4, spaceAfter=4),
        "bullet": ParagraphStyle("bullet", parent=base["Normal"], fontName="Helvetica",
                                 fontSize=10.5, leading=15),
    }

    story = [Paragraph(escape(document["title"]), styles["title"]),
             Paragraph(f"Date: {escape(document['date'])}", styles["meta"])]

    if document["clean"]:
        story.append(Paragraph("No data-quality issues were found in this workbook. "
                               "No corrections are needed.", styles["body"]))
    else:
        for section in document["sections"]:
            story.append(Paragraph(
                escape(f"{section['number']}. {section['title']}"), styles["heading"]))
            story.append(Paragraph(escape(section["problem"]), styles["body"]))
            story.append(Paragraph("Please:", styles["body"]))
            story.append(ListFlowable(
                [ListItem(Paragraph(escape(a), styles["bullet"]), leftIndent=14)
                 for a in section["actions"]],
                bulletType="bullet", bulletFontSize=8, leftIndent=12, spaceAfter=4,
            ))
            for detail in section["details"]:
                story.append(Paragraph(escape(detail["heading"]), styles["note"]))
                items = list(detail["items"])
                if detail["more"]:
                    items.append(f"... and {detail['more']} more")
                if items:
                    story.append(ListFlowable(
                        [ListItem(Paragraph(escape(i), styles["bullet"]), leftIndent=14)
                         for i in items],
                        bulletType="bullet", bulletFontSize=8, leftIndent=12,
                    ))
            story.append(Spacer(1, 4))

    buffer = io.BytesIO()
    SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=document["title"],
    ).build(story)
    return buffer.getvalue()


def to_findings_csv(result) -> bytes:
    """Detailed findings, for the reviewer rather than the data-entry team."""
    import csv

    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    writer.writerow(["Section", "Sheet", "Column", "Issue", "Affected rows",
                     "Excel rows", "Example values"])
    for issue in result.issues:
        rows = sorted(set(issue.rows))
        shown = ", ".join(str(r) for r in rows[:60])
        if len(rows) > 60:
            shown += f", ... (+{len(rows) - 60})"
        writer.writerow([
            issue.section, issue.sheet, issue.column, issue.title,
            issue.count, shown, " | ".join(issue.values[:25]),
        ])
    return buffer.getvalue().encode("utf-8-sig")


def to_preprocess_summary_csv(result) -> bytes:
    """One row per sheet, summarising what the cleaner changed."""
    import csv

    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    writer.writerow(["Sheet", "Rows before", "Rows after", "Blank/placeholder removed",
                     "Duplicates removed", "Date columns", "Dates reformatted",
                     "Dates that could not be read"])
    for s in result.summaries:
        writer.writerow([
            s.sheet, s.original_rows, s.remaining_rows,
            s.removed_blank_or_placeholder, s.removed_duplicates,
            ", ".join(s.date_columns), s.dates_reformatted, len(s.dates_unparseable),
        ])
    return buffer.getvalue().encode("utf-8-sig")


def to_unparseable_dates_csv(result) -> bytes:
    """Every date cell the cleaner left untouched because it could not tell
    what date it represents — for someone to fix by hand."""
    import csv

    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    writer.writerow(["Sheet", "Column", "Excel row", "Original value"])
    for s in result.summaries:
        for u in s.dates_unparseable:
            writer.writerow([u.sheet, u.column, u.row, u.value])
    return buffer.getvalue().encode("utf-8-sig")
